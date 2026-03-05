"""
Generate Mpango ERP v0.2.0 User Manual (Word .docx format).

Usage:
    python scripts/generate_user_manual.py

Output:
    docs/Mpango_ERP_v0.2.0_用户使用手册.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ---------------------------------------------------------------------------
# Output path
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "docs", "Mpango_ERP_v0.2.0_用户使用手册.docx")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex: str):
    """Set background color of a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers: list, rows: list):
    """Add a formatted table with shaded header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, "D9D9D9")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")  # spacing
    return table


def add_screenshot_placeholder(doc, label: str):
    """Insert a styled screenshot placeholder box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"📷  [ 截图：{label} ]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True
    # Add border-like spacing
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(6)
    p_fmt.space_after = Pt(6)


def add_step(doc, step_text: str):
    """Add a numbered-style step paragraph."""
    p = doc.add_paragraph(step_text, style="List Number")
    for run in p.runs:
        run.font.size = Pt(12)


def add_bullet(doc, text: str, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    for run in p.runs:
        run.font.size = Pt(12)


def add_tip(doc, text: str):
    """Add a tip/note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"💡 提示：{text}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x0B, 0x5C, 0xAB)


def add_warning(doc, text: str):
    """Add a warning paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️ 注意：{text}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)


def add_body(doc, text: str):
    """Add a normal body paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = Pt(18)  # ~1.5x for 12pt
    for run in p.runs:
        run.font.size = Pt(12)


# ---------------------------------------------------------------------------
# Configure styles
# ---------------------------------------------------------------------------

def configure_styles(doc):
    """Set up heading and body styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = Pt(18)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    h3.paragraph_format.space_before = Pt(4)
    h3.paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# Document content
# ---------------------------------------------------------------------------

def build_cover_page(doc):
    """Build the cover page."""
    for _ in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Mpango ERP")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("用户使用手册")
    run2.font.size = Pt(24)
    run2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("版本 0.2.0")
    run3.font.size = Pt(16)
    run3.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.add_paragraph("")

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("面向非洲批发–零售供应链的多租户 ERP 系统")
    run4.font.size = Pt(12)
    run4.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

    for _ in range(4):
        doc.add_paragraph("")

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Mpango Team · 2026")
    run5.font.size = Pt(11)
    run5.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)

    doc.add_page_break()


def build_toc_placeholder(doc):
    """Build a table of contents placeholder."""
    doc.add_heading("目录", level=1)
    add_body(doc, "（请在 Word 中右键此处，选择「更新域」以自动生成目录页码。）")

    # Insert TOC field
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)

    doc.add_page_break()


def build_chapter_1(doc):
    """Chapter 1: Introduction."""
    doc.add_heading("第 1 章：简介", level=1)

    doc.add_heading("1.1 Mpango ERP 是什么", level=2)
    add_body(doc,
        "Mpango ERP 是一套专为非洲批发–零售供应链设计的多租户企业资源管理系统。"
        "它帮助肯尼亚及东非地区的批发商（Wholesaler）高效管理库存、销售订单、"
        "应收账款和零售商关系，同时为零售商（Retailer）提供便捷的在线下单和支付体验。"
    )

    doc.add_heading("1.2 系统解决的核心问题", level=2)
    add_bullet(doc, "批发商：手工记账易出错、库存不透明、应收账款难追踪")
    add_bullet(doc, "零售商：电话下单效率低、无法实时查看商品库存和订单状态")
    add_bullet(doc, "供应链：信息断层导致缺货、超卖、对账困难")

    doc.add_heading("1.3 三类用户角色", level=2)

    add_table(doc,
        ["角色", "说明", "典型操作"],
        [
            ["超级管理员\n(Super Admin)", "平台级管理员，管理所有租户（批发商）和系统用户", "创建/停用租户、分配管理员账号、查看审计日志"],
            ["批发商 / 供应商\n(Wholesaler)", "日常运营主角，管理库存、接收订单、处理财务", "商品管理、库存盘点、订单审批、收款确认、报表导出"],
            ["零售商\n(Retailer)", "批发商的下游客户，在线浏览商品并下单", "浏览商品、创建订单、录入支付信息、查看订单状态"],
        ]
    )

    doc.add_heading("1.4 系统架构概览", level=2)
    add_body(doc,
        "Mpango ERP 采用「模块化单体」架构（Modular Monolith），前端使用 React + TypeScript，"
        "后端使用 FastAPI (Python)，数据库为 PostgreSQL，缓存为 Redis。"
        "系统通过 JWT 令牌实现身份认证，通过 RBAC（基于角色的访问控制）实现权限管理，"
        "通过 Schema-per-Tenant 实现多租户数据隔离。"
    )

    doc.add_page_break()


def build_chapter_2(doc):
    """Chapter 2: Quick Start."""
    doc.add_heading("第 2 章：快速开始", level=1)

    doc.add_heading("2.1 系统访问方式", level=2)
    add_body(doc, "在浏览器地址栏输入系统地址即可访问：")
    add_bullet(doc, "生产环境：http://<您的域名或 IP 地址>")
    add_bullet(doc, "演示环境：http://143.110.177.2")

    doc.add_heading("2.2 推荐浏览器", level=2)
    add_table(doc,
        ["浏览器", "最低版本", "推荐"],
        [
            ["Google Chrome", "90+", "✅ 首选"],
            ["Microsoft Edge", "90+", "✅ 推荐"],
            ["Safari", "14+", "✅ 支持"],
            ["Firefox", "88+", "✅ 支持"],
        ]
    )
    add_warning(doc, "不支持 Internet Explorer。建议使用桌面浏览器以获得最佳体验，移动端建议横屏使用。")

    doc.add_heading("2.3 登录系统", level=2)
    add_step(doc, "打开浏览器，输入系统地址。")
    add_step(doc, "在登录页面输入您的邮箱地址和密码。")
    add_step(doc, "点击「登录」按钮。")
    add_step(doc, "登录成功后，系统将自动跳转到首页（Dashboard）。")
    add_screenshot_placeholder(doc, "登录页面")

    add_tip(doc, "如果忘记密码，请联系您的批发商管理员或平台管理员重置密码。")

    doc.add_heading("2.4 首页布局说明", level=2)
    add_body(doc, "登录成功后，您将看到以下界面布局：")
    add_bullet(doc, "左侧边栏：导航菜单，包含 Home（首页）、Sales（销售）、Stock（库存）、Money（财务）、Customers（客户）")
    add_bullet(doc, "顶部栏：面包屑导航 + 当前用户信息")
    add_bullet(doc, "主内容区：当前页面的数据和操作")
    add_screenshot_placeholder(doc, "首页布局总览")

    doc.add_page_break()


def build_chapter_3(doc):
    """Chapter 3: Super Admin Manual."""
    doc.add_heading("第 3 章：超级管理员使用手册", level=1)

    # 3.1
    doc.add_heading("3.1 登录与首页", level=2)
    add_body(doc, "使用超级管理员账号登录系统（例如：admin@mpango.demo / DemoAdmin2026!）。")
    add_body(doc, "登录后，首页将显示平台级概览信息，包括活跃租户数量、系统健康状态等关键指标。")
    add_screenshot_placeholder(doc, "超级管理员首页")

    # 3.2
    doc.add_heading("3.2 租户（批发商）管理", level=2)
    add_body(doc, "超级管理员可以创建、编辑和管理平台上的所有批发商租户。每个租户拥有独立的数据空间，彼此完全隔离。")

    doc.add_heading("3.2.1 查看租户列表", level=3)
    add_step(doc, "点击左侧菜单「Customers」。")
    add_step(doc, "系统显示所有已注册的租户列表，包含名称、编码、状态等信息。")
    add_screenshot_placeholder(doc, "租户列表")

    doc.add_heading("3.2.2 新建租户", level=3)
    add_step(doc, "在租户列表页面，点击右上角「新建租户」按钮。")
    add_step(doc, "填写租户信息：")
    add_bullet(doc, "租户名称（如：Jambo Wholesale Ltd.）")
    add_bullet(doc, "租户编码（如：JAMBO，大写字母+数字，创建后不可修改）")
    add_bullet(doc, "国家（如：Kenya）")
    add_bullet(doc, "默认货币（如：KES）")
    add_step(doc, "点击「保存」，系统自动创建租户的独立数据空间。")
    add_screenshot_placeholder(doc, "新建租户表单")

    doc.add_heading("3.2.3 编辑 / 停用租户", level=3)
    add_body(doc, "在租户列表中，点击某一行的「编辑」按钮可修改租户信息（编码除外）。")
    add_warning(doc, "停用租户后，该租户下的所有用户将无法登录，但数据不会被删除。")

    # 3.3
    doc.add_heading("3.3 用户与角色管理", level=2)
    add_body(doc, "超级管理员可以为每个租户创建管理员账号和运营人员账号，并分配不同的角色。")

    doc.add_heading("3.3.1 创建租户管理员", level=3)
    add_step(doc, "进入目标租户的详情页。")
    add_step(doc, "点击「用户管理」标签页。")
    add_step(doc, "点击「新建用户」，填写邮箱、姓名，选择角色为「Admin」。")
    add_step(doc, "点击「保存」，系统将发送初始密码至该邮箱。")
    add_screenshot_placeholder(doc, "新建用户页面")

    doc.add_heading("3.3.2 角色与权限说明", level=3)
    add_table(doc,
        ["角色", "可访问模块", "关键权限"],
        [
            ["Admin（管理员）", "全部模块", "用户管理、商品管理、订单管理、财务管理、报表导出"],
            ["Sales（销售）", "Home, Sales", "查看/创建/确认订单、查看客户信息"],
            ["Warehouse（仓管）", "Home, Stock", "库存查看、入库/出库操作、盘点"],
            ["Finance（财务）", "Home, Sales(只读), Money", "收款确认、发票管理、财务报表"],
        ]
    )
    add_screenshot_placeholder(doc, "用户列表")

    # 3.4
    doc.add_heading("3.4 审计与系统监控", level=2)
    add_body(doc, "系统自动记录所有关键操作的审计日志，包括：谁（操作人）、在何时（时间戳）、做了什么（操作类型）、影响了哪些数据。")
    add_bullet(doc, "常见用途：安全审查、排查误操作、合规审计")
    add_bullet(doc, "日志保留策略：所有日志永久保留，不可删除")
    add_screenshot_placeholder(doc, "审计日志页面")

    doc.add_page_break()


def build_chapter_4(doc):
    """Chapter 4: Wholesaler Manual."""
    doc.add_heading("第 4 章：批发商 / 供应商使用手册", level=1)

    # 4.1
    doc.add_heading("4.1 首次登录与基础设置", level=2)
    add_body(doc, "当超级管理员为您创建账号后，您将收到一封包含初始密码的邮件。")
    add_step(doc, "打开系统登录页面，输入您的邮箱和初始密码。")
    add_step(doc, "首次登录后，系统可能引导您完成基础设置（Onboarding）：")
    add_bullet(doc, "确认公司信息（名称、地址、联系方式）")
    add_bullet(doc, "设置默认货币（如 KES）")
    add_bullet(doc, "配置主仓库信息")
    add_step(doc, "完成设置后，进入系统首页。")
    add_screenshot_placeholder(doc, "Onboarding 向导步骤 1")
    add_screenshot_placeholder(doc, "Onboarding 向导步骤 2")

    # 4.2
    doc.add_heading("4.2 仓库与供应商资料", level=2)

    doc.add_heading("4.2.1 管理仓库", level=3)
    add_body(doc, "作为批发商，您可以管理一个或多个仓库。每个仓库有独立的库存记录。")
    add_step(doc, "点击左侧菜单「Stock」。")
    add_step(doc, "查看当前仓库列表。")
    add_step(doc, "如需新建仓库，点击「新建仓库」，填写名称、地址、负责人。")
    add_screenshot_placeholder(doc, "仓库列表")

    doc.add_heading("4.2.2 维护供应商信息", level=3)
    add_body(doc, "记录您的上游供应商信息，便于采购管理和对账。")
    add_bullet(doc, "供应商名称、联系人、电话、邮箱")
    add_bullet(doc, "支付方式（现金 / 银行转账 / M-Pesa / 账期）")
    add_bullet(doc, "供应品类")

    # 4.3
    doc.add_heading("4.3 商品管理", level=2)

    doc.add_heading("4.3.1 查看商品列表", level=3)
    add_step(doc, "点击左侧菜单「Stock」，进入库存页面。")
    add_step(doc, "系统显示所有商品的 SKU 编码、名称、库存数量、可用数量等信息。")
    add_step(doc, "使用搜索框可按商品名称或 SKU 快速查找。")
    add_screenshot_placeholder(doc, "商品列表页")

    doc.add_heading("4.3.2 新建商品", level=3)
    add_step(doc, "在商品列表页，点击右上角「新建商品」按钮。")
    add_step(doc, "填写商品信息：")
    add_bullet(doc, "商品名称（如：Unga Maize Meal 2kg）")
    add_bullet(doc, "SKU 编码（如：P001，系统唯一）")
    add_bullet(doc, "单位（如：Bale、Carton、Jerrycan）")
    add_bullet(doc, "分类（如：粮食、食用油、调味料、日用品、饮料）")
    add_bullet(doc, "单价（KES）")
    add_step(doc, "点击「保存」，商品将出现在列表中。")
    add_screenshot_placeholder(doc, "新建商品页面")

    add_tip(doc,
        "v0.2.0 暂不支持商品图片上传功能。如需为商品添加图片，请参阅本手册「附录 A：v0.2.1 UI 改进建议」。"
    )

    # 4.4
    doc.add_heading("4.4 库存管理", level=2)
    add_body(doc, "库存页面实时显示每个 SKU 的库存状态：")
    add_table(doc,
        ["字段", "说明"],
        [
            ["库存总量 (On Hand)", "仓库中的实际数量"],
            ["已预留 (Reserved)", "已被订单锁定但尚未发货的数量"],
            ["可用数量 (Available)", "= 库存总量 - 已预留，可供新订单使用"],
        ]
    )
    add_screenshot_placeholder(doc, "库存列表（含库存状态）")

    add_warning(doc, "当零售商下单时，系统会自动检查可用库存。如果库存不足，订单将无法创建。")

    # 4.5
    doc.add_heading("4.5 接收零售商订单与发货", level=2)

    doc.add_heading("4.5.1 查看新订单", level=3)
    add_step(doc, "点击左侧菜单「Sales」，进入订单列表页。")
    add_step(doc, "新订单将显示在列表顶部，状态为「Draft」或「Confirmed」。")
    add_step(doc, "点击订单行可查看订单详情，包括零售商信息、商品明细、金额。")
    add_screenshot_placeholder(doc, "订单列表")

    doc.add_heading("4.5.2 订单状态流转", level=3)
    add_body(doc, "Mpango ERP 使用严格的订单状态机，确保每笔交易可追溯：")
    add_table(doc,
        ["状态", "英文", "说明", "可执行操作"],
        [
            ["草稿", "Draft", "零售商刚创建的订单", "确认 / 取消"],
            ["已确认", "Confirmed", "批发商已审核通过", "标记已付款 / 取消"],
            ["已付款", "Paid", "已收到零售商付款", "标记已发货"],
            ["已发货", "Fulfilled", "商品已送达零售商", "（终态）"],
            ["已取消", "Cancelled", "订单被取消", "（终态）"],
            ["已退货", "Returned", "零售商发起退货", "（终态）"],
        ]
    )

    doc.add_heading("4.5.3 确认订单", level=3)
    add_step(doc, "在订单详情页，检查商品明细和库存是否充足。")
    add_step(doc, "点击「确认订单」按钮，订单状态变为「Confirmed」。")
    add_step(doc, "系统自动预留对应库存（Reserved 数量增加）。")

    doc.add_heading("4.5.4 记录付款", level=3)
    add_step(doc, "收到零售商付款后，在订单详情页点击「标记已付款」。")
    add_step(doc, "填写支付信息：")
    add_bullet(doc, "支付方式：M-Pesa / 银行转账 / 现金")
    add_bullet(doc, "交易参考号（如 M-Pesa 参考号：QH8923XZLP）")
    add_step(doc, "点击「确认」，订单状态变为「Paid」。")

    doc.add_heading("4.5.5 发货完成", level=3)
    add_step(doc, "商品发出后，在订单详情页点击「标记已发货」。")
    add_step(doc, "订单状态变为「Fulfilled」，库存正式扣减。")
    add_screenshot_placeholder(doc, "订单详情页（含状态操作按钮）")

    doc.add_heading("4.5.6 测试场景示例", level=3)
    add_body(doc, "场景 A：正常订单流程")
    add_bullet(doc, "零售商 Grace Wanjiku (Mama Mboga Shop) 下单 5 箱 Unga Maize Meal")
    add_bullet(doc, "批发商审批 → 确认订单 → 库存预留 5 箱")
    add_bullet(doc, "Grace 通过 M-Pesa 付款（参考号：QH8923XZLP）")
    add_bullet(doc, "批发商标记已付款 → 发货 → 订单完成，库存正式扣减")

    add_body(doc, "场景 B：缺货订单")
    add_bullet(doc, "零售商 John Kamau (Kiosk 254) 尝试下单 10 箱 Soko Maize Meal (P004)")
    add_bullet(doc, "系统提示：「库存不足，无法下单」（P004 库存为 0）")

    # 4.6
    doc.add_heading("4.6 财务与应收账款", level=2)
    add_body(doc, "「Money」模块帮助您管理应收账款（Accounts Receivable），追踪每个零售商的未结清款项。")

    doc.add_heading("4.6.1 查看应收账款", level=3)
    add_step(doc, "点击左侧菜单「Money」。")
    add_step(doc, "系统显示所有零售商的应收账款汇总：已开票金额、已收款金额、未结清余额。")
    add_screenshot_placeholder(doc, "应收账款列表页")

    doc.add_heading("4.6.2 导出报表", level=3)
    add_step(doc, "在财务页面，点击「导出 CSV」按钮。")
    add_step(doc, "系统生成包含所有交易记录的 CSV 文件，可用 Excel 打开进行进一步分析。")
    add_screenshot_placeholder(doc, "导出按钮与 CSV 下载")

    doc.add_page_break()


def build_chapter_5(doc):
    """Chapter 5: Retailer Manual."""
    doc.add_heading("第 5 章：零售商使用手册", level=1)

    # 5.1
    doc.add_heading("5.1 注册 / 接受邀请", level=2)
    add_body(doc, "零售商账号由批发商管理员创建。创建后，您将收到一封包含登录信息的邮件。")
    add_step(doc, "查收邮件，获取您的登录邮箱和初始密码。")
    add_step(doc, "打开系统登录页面，输入邮箱和密码。")
    add_step(doc, "首次登录后，建议立即修改密码。")
    add_screenshot_placeholder(doc, "零售商邀请邮件示例")
    add_screenshot_placeholder(doc, "零售商登录页")

    # 5.2
    doc.add_heading("5.2 浏览商品与创建订单", level=2)

    doc.add_heading("5.2.1 浏览商品", level=3)
    add_step(doc, "登录后，点击左侧菜单「Stock」或首页的商品入口。")
    add_step(doc, "系统显示批发商提供的所有商品列表，包含名称、SKU、单价、可用库存。")
    add_step(doc, "使用搜索框或分类过滤器快速找到目标商品。")
    add_screenshot_placeholder(doc, "商品列表（零售商视角）")

    add_tip(doc,
        "v0.2.0 暂不支持商品图片显示。您可通过 SKU 编码和商品名称识别商品。"
        "图片功能将在 v0.2.1 中上线。"
    )

    doc.add_heading("5.2.2 创建订单", level=3)
    add_step(doc, "点击左侧菜单「Sales」，然后点击「新建订单」。")
    add_step(doc, "在订单页面，添加商品：")
    add_bullet(doc, "搜索或选择商品")
    add_bullet(doc, "输入订购数量")
    add_bullet(doc, "系统自动计算小计和总金额")
    add_step(doc, "确认商品和数量无误后，点击「提交订单」。")
    add_step(doc, "订单提交成功，状态为「Draft」，等待批发商确认。")
    add_screenshot_placeholder(doc, "新建订单页")
    add_screenshot_placeholder(doc, "订单确认页")

    add_warning(doc,
        "如果某商品库存为 0 或库存不足，系统将提示「库存不足，无法下单」。"
        "请联系批发商确认补货时间。"
    )

    # 5.3
    doc.add_heading("5.3 支付与查看订单状态", level=2)

    doc.add_heading("5.3.1 查看订单状态", level=3)
    add_step(doc, "点击左侧菜单「Sales」，查看您的所有订单。")
    add_step(doc, "每笔订单显示：订单编号、创建时间、商品数量、总金额、当前状态。")
    add_step(doc, "点击订单行可查看详细信息。")
    add_screenshot_placeholder(doc, "我的订单列表")

    doc.add_heading("5.3.2 录入支付信息", level=3)
    add_body(doc, "当批发商确认订单后（状态变为「Confirmed」），您需要完成付款：")
    add_step(doc, "通过 M-Pesa 或银行转账向批发商付款。")
    add_step(doc, "在订单详情页，点击「录入支付信息」。")
    add_step(doc, "填写：")
    add_bullet(doc, "支付方式（M-Pesa / 银行转账）")
    add_bullet(doc, "交易参考号（如 M-Pesa 参考号：QH8923XZLP）")
    add_step(doc, "点击「提交」，等待批发商确认收款。")
    add_screenshot_placeholder(doc, "订单详情（含支付信息）")

    # 5.4
    doc.add_heading("5.4 退货与问题反馈", level=2)

    doc.add_heading("5.4.1 发起退货", level=3)
    add_body(doc, "如果收到的商品存在质量问题或与订单不符，您可以发起退货请求。")
    add_step(doc, "在订单详情页，点击「退货」按钮。")
    add_step(doc, "填写退货原因（如：包装破损、商品过期、数量不符）。")
    add_step(doc, "提交后，等待批发商审核。")
    add_step(doc, "退货审核通过后，订单状态变为「Returned」。")

    doc.add_heading("5.4.2 常见问题联系方式", level=3)
    add_bullet(doc, "登录问题：联系批发商管理员重置密码")
    add_bullet(doc, "权限问题：联系批发商管理员调整角色")
    add_bullet(doc, "系统故障：联系平台技术支持")

    doc.add_page_break()


def build_chapter_6(doc):
    """Chapter 6: FAQ."""
    doc.add_heading("第 6 章：常见问题（FAQ）", level=1)

    # Q1
    doc.add_heading("Q1：我看不到某些菜单，是不是系统出错？", level=2)
    add_body(doc,
        "不是系统故障。Mpango ERP 使用基于角色的访问控制（RBAC），不同角色只能看到自己有权限的模块。"
        "例如，仓管人员只能看到「Home」和「Stock」，无法看到「Money」（财务）模块。"
        "如果您需要访问更多模块，请联系管理员调整您的角色权限。"
    )

    # Q2
    doc.add_heading("Q2：为何同一账号不能看到其他城市的仓库？", level=2)
    add_body(doc,
        "这是多租户隔离机制（Tenant Guardrail）的设计。每个批发商是一个独立的租户，"
        "拥有自己的数据空间。不同租户之间的数据完全隔离，即使是同一个人拥有多个批发商账号，"
        "也需要分别登录对应的租户才能查看各自的数据。这是为了保护商业数据安全。"
    )

    # Q3
    doc.add_heading("Q3：忘记密码怎么办？", level=2)
    add_step(doc, "联系您的批发商管理员或平台超级管理员。")
    add_step(doc, "管理员可以在「用户管理」中为您重置密码。")
    add_step(doc, "您将收到一封包含新密码的邮件。")
    add_tip(doc, "v0.2.0 暂不支持自助找回密码功能。该功能计划在后续版本中上线。")

    # Q4
    doc.add_heading("Q4：手机上显示不完整怎么办？", level=2)
    add_body(doc,
        "Mpango ERP v0.2.0 主要针对桌面浏览器优化。如果您在手机上使用，建议："
    )
    add_bullet(doc, "将手机横屏使用")
    add_bullet(doc, "使用平板电脑")
    add_bullet(doc, "优先使用桌面电脑或笔记本电脑")
    add_tip(doc, "移动端适配（PWA）计划在 Phase 3 中实现。")

    # Q5
    doc.add_heading("Q5：商品列表没有显示图片？", level=2)
    add_body(doc,
        "这是 v0.2.0 的已知限制。当前版本暂不支持商品图片上传和显示功能。"
        "您可以通过 SKU 编码和商品名称来识别商品。"
        "商品图片功能将在 v0.2.1 版本中上线，届时商品列表和订单详情页将显示商品缩略图。"
    )

    # Q6
    doc.add_heading("Q6：导出的 CSV 文件用 Excel 打开是乱码？", level=2)
    add_body(doc, "这通常是编码问题。解决方法：")
    add_step(doc, "打开 Excel，选择「数据」→「从文本/CSV」。")
    add_step(doc, "选择下载的 CSV 文件。")
    add_step(doc, "在导入向导中，将编码设置为「UTF-8」。")
    add_step(doc, "点击「加载」即可正常显示。")

    doc.add_page_break()


def build_appendix_a(doc):
    """Appendix A: v0.2.1 UI Improvement Suggestions."""
    doc.add_heading("附录 A：v0.2.1 UI 改进建议（商品图片功能）", level=1)

    p = doc.add_paragraph()
    run = p.add_run(
        "⚠️ 当前限制：Mpango ERP v0.2.0 的商品管理模块暂时不支持图片上传和显示功能。"
        "这会影响零售商在浏览商品时的体验，尤其是快速识别商品品牌和包装。"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

    doc.add_heading("建议改进内容（v0.2.1 优先级：高）", level=2)

    add_table(doc,
        ["序号", "改进项", "说明"],
        [
            ["1", "数据库层：增加 image_url 字段",
             "在 Product / SKU 实体中增加 image_url 字段（VARCHAR(512)），存储图片链接"],
            ["2", "前端商品列表页：显示缩略图",
             "在每行左侧显示 80×80px 缩略图，无图片时显示默认占位图"],
            ["3", "前端订单详情页：显示商品图片",
             "在订单行项目中显示商品缩略图，帮助仓管核对发货"],
            ["4", "商品编辑页：图片上传功能",
             "支持本地上传至云存储（AWS S3 / Cloudflare R2），自动生成缩略图"],
            ["5", "默认占位图",
             "为无图片商品显示品类图标（粮食、饮料、日用品等）"],
        ]
    )

    doc.add_heading("业务价值", level=2)
    add_bullet(doc, "提升零售商下单效率（视觉识别比纯文字快 3 倍）")
    add_bullet(doc, "减少因商品混淆导致的退换货率（预计降低 15-20%）")
    add_bullet(doc, "增强系统专业度，符合现代电商体验标准")
    add_bullet(doc, "为未来移动端 App / PWA 打下基础")

    doc.add_heading("技术实现路径", level=2)
    add_step(doc, "后端：在 SKU 模型中增加 image_url: Optional[str] 字段，更新 Alembic 迁移")
    add_step(doc, "后端：新增 /api/v1/upload 端点，接收图片文件，上传至 S3/R2，返回 URL")
    add_step(doc, "前端：在 StockView 类型中增加 image_url 字段")
    add_step(doc, "前端：在 InventoryPage 和 OrderListPage 中渲染缩略图组件")
    add_step(doc, "前端：在商品编辑表单中增加图片上传控件")

    doc.add_page_break()


def build_appendix_b(doc):
    """Appendix B: Test accounts quick reference."""
    doc.add_heading("附录 B：测试账号速查表", level=1)

    add_table(doc,
        ["角色", "邮箱", "密码", "可访问模块"],
        [
            ["超级管理员", "admin@mpango.demo", "DemoAdmin2026!", "全部模块 + 租户管理"],
            ["批发商管理员", "admin@jambo.co.ke", "Password123!", "Home, Sales, Stock, Money, Customers"],
            ["批发商仓管", "warehouse@jambo.co.ke", "Warehouse123!", "Home, Stock"],
            ["批发商财务", "finance@jambo.co.ke", "Finance123!", "Home, Sales(只读), Money"],
            ["零售商 A", "grace@mamammboga.co.ke", "Retail123!", "商品浏览, 下单, 我的订单"],
            ["零售商 B", "john@kiosk254.co.ke", "Retail123!", "商品浏览, 下单, 我的订单"],
            ["零售商 C", "amina@westlandsminimart.co.ke", "Retail123!", "商品浏览, 下单, 我的订单"],
            ["零售商 D", "mohamed@eastleighduka.co.ke", "Retail123!", "商品浏览, 下单, 我的订单"],
        ]
    )

    add_body(doc, "")
    add_body(doc, "— 文档结束 —")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document()
    configure_styles(doc)

    # Build all sections
    build_cover_page(doc)
    build_toc_placeholder(doc)
    build_chapter_1(doc)
    build_chapter_2(doc)
    build_chapter_3(doc)
    build_chapter_4(doc)
    build_chapter_5(doc)
    build_chapter_6(doc)
    build_appendix_a(doc)
    build_appendix_b(doc)

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"✅ User manual generated: {OUTPUT_PATH}")
    print(f"   File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")


if __name__ == "__main__":
    main()
